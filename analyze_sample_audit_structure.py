"""Analyze sample audit DOCX files to understand section depth and structure."""
from docx import Document
import sys

def analyze_docx(filepath):
    """Analyze a DOCX file and return section structure."""
    doc = Document(filepath)
    
    sections = {}
    current_section = None
    para_count = 0
    table_count = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Detect section headers (Heading 1 or all caps short text)
        is_header = (
            para.style.name.startswith('Heading 1') or 
            (len(text) < 100 and text.isupper() and len(text) > 5) or
            para.style.name.startswith('Heading 2')
        )
        
        if is_header:
            current_section = text[:80]  # Limit length
            if current_section not in sections:
                sections[current_section] = {
                    'paragraphs': 0,
                    'tables': 0,
                    'first_para': None
                }
        
        if current_section:
            sections[current_section]['paragraphs'] += 1
            if sections[current_section]['first_para'] is None and len(text) > 20:
                sections[current_section]['first_para'] = text[:100]
    
    # Count tables per section (approximate - tables are global in docx)
    table_count = len(doc.tables)
    
    return sections, len(doc.paragraphs), table_count

if __name__ == "__main__":
    filepath = "SampleAudit/Audit Example - Clothing & Accessories.docx"
    
    try:
        sections, total_paras, total_tables = analyze_docx(filepath)
        
        print(f"\n📊 ANALYSIS: {filepath}")
        print(f"Total Paragraphs: {total_paras}")
        print(f"Total Tables: {total_tables}")
        print(f"Total Sections: {len(sections)}\n")
        
        print("=" * 80)
        print("SECTION BREAKDOWN:")
        print("=" * 80)
        
        for i, (section_name, data) in enumerate(sections.items(), 1):
            print(f"\n{i}. {section_name}")
            print(f"   Paragraphs: {data['paragraphs']}")
            if data['first_para']:
                print(f"   Sample: {data['first_para']}...")
        
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

