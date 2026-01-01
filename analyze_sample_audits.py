"""Analyze sample audit Word documents to understand structure and content."""
from docx import Document
import os
from collections import defaultdict

def analyze_docx(filepath):
    """Analyze a Word document and extract key information."""
    doc = Document(filepath)
    
    analysis = {
        'file': os.path.basename(filepath),
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'images': len([r for r in doc.part.rels.values() if 'image' in r.target_ref]),
        'sections': [],
        'tables_info': [],
        'has_charts': False,
        'key_content': []
    }
    
    # Extract sections (headings/titles)
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        # Detect section headers (uppercase, short, or contains key words)
        if (text.isupper() and len(text) < 100) or \
           any(word in text for word in ['KAV', 'Campaign', 'Flow', 'List', 'Form', 'Strategic', 'Recommendation', 'Section']):
            analysis['sections'].append(text)
        
        # Check for chart references
        if any(word in text.lower() for word in ['chart', 'graph', 'visual', 'diagram', 'figure', 'line graph', 'bar chart']):
            analysis['has_charts'] = True
            analysis['key_content'].append(f"Chart reference: {text[:150]}")
    
    # Analyze tables
    for i, table in enumerate(doc.tables[:10]):
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            analysis['tables_info'].append({
                'index': i + 1,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'headers': headers[:5]  # First 5 headers
            })
    
    return analysis

# Analyze all sample audits
print("=" * 80)
print("SAMPLE AUDIT ANALYSIS")
print("=" * 80)

sample_files = [
    'SampleAudit/Audit Process as of 1_1_26.docx',
    'SampleAudit/Audit Example - Clothing & Accessories.docx',
    'SampleAudit/Audit Example - Food & Beverage.docx',
    'SampleAudit/Audit Example - Health & Beauty.docx',
    'SampleAudit/Audit Example - Home.docx',
    'SampleAudit/Audit Example - Specialty.docx'
]

all_analyses = {}
for filepath in sample_files:
    if os.path.exists(filepath):
        print(f"\n{'='*80}")
        print(f"Analyzing: {os.path.basename(filepath)}")
        print('='*80)
        analysis = analyze_docx(filepath)
        all_analyses[filepath] = analysis
        
        print(f"Paragraphs: {analysis['paragraphs']}")
        print(f"Tables: {analysis['tables']}")
        print(f"Images: {analysis['images']}")
        print(f"Has Charts: {analysis['has_charts']}")
        
        print(f"\nSections Found ({len(analysis['sections'])}):")
        for i, section in enumerate(analysis['sections'][:20], 1):
            print(f"  {i}. {section[:100]}")
        
        if analysis['tables_info']:
            print(f"\nTables Structure:")
            for table_info in analysis['tables_info'][:5]:
                print(f"  Table {table_info['index']}: {table_info['rows']}x{table_info['cols']}")
                print(f"    Headers: {table_info['headers']}")
        
        if analysis['key_content']:
            print(f"\nKey Content:")
            for content in analysis['key_content'][:5]:
                print(f"  - {content}")

# Compare structure across audits
print(f"\n{'='*80}")
print("COMPARATIVE ANALYSIS")
print('='*80)

industry_audits = [f for f in sample_files if 'Example' in f]
if industry_audits:
    print(f"\nAnalyzing {len(industry_audits)} industry-specific audits...")
    
    common_sections = defaultdict(int)
    for filepath in industry_audits:
        if filepath in all_analyses:
            for section in all_analyses[filepath]['sections']:
                # Normalize section names
                section_lower = section.lower()
                if 'kav' in section_lower:
                    common_sections['KAV'] += 1
                elif 'campaign' in section_lower:
                    common_sections['Campaigns'] += 1
                elif 'flow' in section_lower:
                    common_sections['Flows'] += 1
                elif 'list' in section_lower or 'database' in section_lower:
                    common_sections['List/Database'] += 1
                elif 'form' in section_lower or 'data capture' in section_lower:
                    common_sections['Forms/Data Capture'] += 1
                elif 'strategic' in section_lower or 'recommendation' in section_lower:
                    common_sections['Strategic Recommendations'] += 1
    
    print("\nCommon Sections Across All Audits:")
    for section, count in sorted(common_sections.items(), key=lambda x: -x[1]):
        print(f"  {section}: {count}/{len(industry_audits)} audits")

