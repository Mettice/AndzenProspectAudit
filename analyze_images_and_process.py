"""Analyze images and audit process in detail."""
from docx import Document
import os

def extract_full_process():
    """Extract the complete audit process."""
    doc = Document('SampleAudit/Audit Process as of 1_1_26.docx')
    
    print("=" * 80)
    print("COMPLETE AUDIT PROCESS DOCUMENT")
    print("=" * 80)
    
    current_section = None
    process_steps = {}
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        # Detect section headers
        if text.startswith('Section') or (text[0].isdigit() and ':' in text):
            current_section = text
            process_steps[current_section] = []
        elif current_section:
            process_steps[current_section].append(text)
        else:
            # Before sections start
            if 'Section' not in str(process_steps):
                if 'steps' not in process_steps:
                    process_steps['steps'] = []
                process_steps['steps'].append(text)
    
    # Print organized process
    if 'steps' in process_steps:
        print("\n--- PRE-AUDIT STEPS ---")
        for step in process_steps['steps']:
            print(f"  • {step}")
    
    for section, steps in process_steps.items():
        if section == 'steps':
            continue
        print(f"\n--- {section} ---")
        for i, step in enumerate(steps, 1):
            print(f"  {i}. {step}")
    
    return process_steps

def analyze_image_types():
    """Analyze what types of images are in sample audits."""
    print("\n" + "=" * 80)
    print("IMAGE TYPE ANALYSIS")
    print("=" * 80)
    
    sample_files = [
        'SampleAudit/Audit Example - Clothing & Accessories.docx',
        'SampleAudit/Audit Example - Food & Beverage.docx',
        'SampleAudit/Audit Example - Health & Beauty.docx',
    ]
    
    image_types = {
        'screenshots': [],
        'charts': [],
        'tables': [],
        'campaign_images': [],
        'other': []
    }
    
    for filepath in sample_files:
        if not os.path.exists(filepath):
            continue
        
        doc = Document(filepath)
        print(f"\n--- {os.path.basename(filepath)} ---")
        
        # Check paragraphs for image context
        for p in doc.paragraphs:
            text = p.text.lower()
            
            if 'screenshot' in text or 'dashboard' in text:
                image_types['screenshots'].append(p.text[:150])
            elif 'chart' in text or 'graph' in text or 'line graph' in text:
                image_types['charts'].append(p.text[:150])
            elif 'table' in text and ('screenshot' in text or 'image' in text):
                image_types['tables'].append(p.text[:150])
            elif 'campaign' in text and ('image' in text or 'creative' in text):
                image_types['campaign_images'].append(p.text[:150])
            elif any(word in text for word in ['image', 'figure', 'visual', 'diagram']):
                if p.text not in [x[:150] for x in image_types['other']]:
                    image_types['other'].append(p.text[:150])
    
    print("\n=== IMAGE TYPE SUMMARY ===")
    print(f"Screenshots: {len(image_types['screenshots'])} references")
    print(f"Charts/Graphs: {len(image_types['charts'])} references")
    print(f"Table Screenshots: {len(image_types['tables'])} references")
    print(f"Campaign Images: {len(image_types['campaign_images'])} references")
    print(f"Other: {len(image_types['other'])} references")
    
    print("\n--- Sample Screenshot References ---")
    for ref in image_types['screenshots'][:5]:
        print(f"  • {ref}...")
    
    print("\n--- Sample Chart References ---")
    for ref in image_types['charts'][:5]:
        print(f"  • {ref}...")
    
    return image_types

if __name__ == '__main__':
    process = extract_full_process()
    image_types = analyze_image_types()
    
    print("\n" + "=" * 80)
    print("RECOMMENDATIONS")
    print("=" * 80)
    print("\n1. CHARTS: We can generate these from our data using Chart.js")
    print("   - Engagement breakdown line charts")
    print("   - Flow performance comparison charts")
    print("   - Revenue trend charts")
    print("\n2. SCREENSHOTS: These are Klaviyo dashboard screenshots")
    print("   - We don't have direct access to capture these")
    print("   - Could be optional or use placeholders")
    print("\n3. CAMPAIGN IMAGES: These are actual email/campaign creatives")
    print("   - We don't have access to these unless pre-populated")
    print("   - Can be skipped or made optional")
    print("\n4. TABLE SCREENSHOTS: These are screenshots of Klaviyo tables")
    print("   - We generate our own tables, so screenshots not needed")
    print("   - Our tables are sufficient")

